import docx
import random

# Eski belgeyi aç
eski_belge = docx.Document("tefsir özet.docx")

# Kullanılabilir fontlar
kullanilabilir_fontlar = ["c11", "c12", "c13", "c14", "c15", "c16", "c17", "c18", "c19", "c20"]

# Yeni bir belge oluştur
yeni_belge = docx.Document()

# Eski belgedeki her paragrafı gez
for eski_paragraf in eski_belge.paragraphs:
    # Paragrafı al
    metin = eski_paragraf.text
    
    # Yeni bir paragraf oluştur
    yeni_paragraf = yeni_belge.add_paragraph()
    
    # Her harf için ayrı bir font seçerek metni ekle
    for harf in metin:
        # Rastgele bir font seç
        rastgele_font = random.choice(kullanilabilir_fontlar)
        
        # Yeni bir run (metin parçası) oluştur ve harfi ekleyerek fontu ayarla
        yeni_run = yeni_paragraf.add_run(harf)
        yeni_run.font.name = rastgele_font

# Yeni belgeyi kaydet
yeni_belge.save("harf_harf_fontlar.docx")
